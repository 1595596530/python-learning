from docx import Document
from docx.shared import Inches
import pyttsx3 

def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile picture
document.add_picture('yasuo.jpg',width=Inches(2.0))
 


#name phone number and email details
name = input('what is your name?')
speak('hello'+ name + ' how are you today ')

speak('what is your phone number?')
phone_number = input('what is your phone number?')
email = input('what is your email?')

document.add_paragraph( 
name + ' | ' + phone_number + ' | ' + email)


# about me
document.add_heading('about me')
document.add_paragraph( 
    input('tell me about yourself? ')
)

#work experience
document.add_heading('work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('from Date ')
to_date = input('To Date ')

p.add_run(company +' ').bold = True
p.add_run(from_date + '-' + to_date +'\n').italic = True

experience_details = input(
    'Decride you experience at ' +company + ' ')
p.add_run(experience_details)

#more experiences
while True:
    has_more_experiences =input(
        'Do you have more experiences? yes or no')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('from Date ')
        to_date = input('To Date ')

        p.add_run(company +' ').bold = True
        p.add_run(from_date + '-' + to_date +'\n').italic = True

        experience_details = input(
            'Decride you experience at ' +company)
        p.add_run(experience_details)
    else:
        break
    
#skill
document.add_heading('skills')
skill= input('Enter skill')
p = document.add_paragraph(skill)


while True:
    has_more_skill = input('dp you has more skill? yes or no')
    if has_more_skill.lower() == 'yes':
        skill=input('Enter skill')
        p = document.add_paragraph(skill)
        p.style = 'list Bullet'
    else:
        break

#footer
section = document.sections[0]
footer = section.footer       
p = footer.paragraphs[0]
p.text = 'CV generated using amigoscode'

document.save('cv.docx')